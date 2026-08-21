# tests/test_regressions.py

"""
Guards for document invariants that are cheap to break.

Every check answers one question: does the DOCX that comes out hold
everything that went in, styled and numbered the way it was asked for?
Stdlib `unittest` plus `python-docx`, which is a base dependency.

  >>> python -m pytest tests/test_regressions.py
"""

import os, tempfile, unittest, warnings

import docx
from docx.oxml.ns import qn

from docmarq import DOCX
from docmarq.md import md_to_docx
from conftest import (docx_text, docx_numbered_paras, docx_start_overrides,
  docx_ppr_children)

class TempDocxCase(unittest.TestCase):
  """Base case giving each test a private directory to write documents into."""
  def setUp(self):
    self._tmp = tempfile.TemporaryDirectory()

  def tearDown(self):
    self._tmp.cleanup()

  def path(self, name:str="out.docx") -> str:
    return os.path.join(self._tmp.name, name)

  def write_svg(self, name:str="tiny.svg") -> str:
    p = os.path.join(self._tmp.name, name)
    with open(p, "w", encoding="utf-8") as f:
      f.write('<svg xmlns="http://www.w3.org/2000/svg" width="40" height="20">'
        '<rect width="40" height="20" fill="#37f"/></svg>')
    return p

  @staticmethod
  def md_url(path:str) -> str:
    """Path spelled for a markdown link target.

    A backslash escapes the next character in CommonMark, so a Windows path
    has to travel with forward slashes.
    """
    return path.replace(os.sep, "/")

#----------------------------------------------------------------------------- Blocks in list items

class TestBlocksInListItems(TempDocxCase):
  """Code, quotes and tables under a bullet are the ordinary shape of technical
  markdown. A list item carries the same content a document body can."""

  MD = (
    "- item one\n\n"
    "  ```python\n  CODEMARKER = 1\n  ```\n\n"
    "- item two\n\n"
    "  > QUOTEMARKER inside list\n\n"
    "- item three\n\n"
    "  | a | b |\n  |---|---|\n  | TABLEMARKER | x |\n"
  )

  def test_every_block_reaches_the_document(self):
    p = self.path()
    md_to_docx(self.MD, p)
    body = docx_text(p)
    for marker in ("CODEMARKER", "QUOTEMARKER", "TABLEMARKER", "item one"):
      self.assertIn(marker, body)

  def test_math_block_in_list_item(self):
    p = self.path()
    md_to_docx("- with math\n\n  $$\n  E = mc^2\n  $$\n", p)
    self.assertIn("oMath", docx.Document(p).element.xml)

#------------------------------------------------------------------------------------------- Images

class TestSvgImages(TempDocxCase):
  """python-docx reads raster formats only, so an SVG has to be rasterized on
  the way in - for the fluent API as much as for markdown figures."""

  def test_fluent_image_accepts_svg(self):
    p = self.path()
    d = DOCX(p)
    d.image(self.write_svg(), width=40)
    d.save()
    self.assertTrue(os.path.exists(p))
    self.assertIn("graphic", docx.Document(p).element.xml)

  def test_markdown_figure_accepts_svg(self):
    p = self.path()
    svg = self.write_svg()
    md_to_docx(f"Before.\n\n![diagram]({self.md_url(svg)})\n\nTAILMARK.\n", p)
    self.assertIn("TAILMARK", docx_text(p))
    self.assertIn("graphic", docx.Document(p).element.xml)

  def test_missing_rasterizer_costs_one_figure(self):
    """No rasterizing backend is a degraded figure, not a dead render.

    `rlPyCairo` needs a native cairo build, so the import can fail on a
    machine that has the package installed.
    """
    import builtins
    blocked = ("rlPyCairo", "svglib", "svglib.svglib")
    real_import = builtins.__import__

    def blocking_import(name, *a, **kw):
      if name in blocked:
        raise ImportError(f"No module named {name!r}")
      return real_import(name, *a, **kw)

    p = self.path()
    svg = self.write_svg()
    builtins.__import__ = blocking_import
    try:
      with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        md_to_docx(f"![ALTMARK]({self.md_url(svg)})\n\nTAILMARK.\n", p)
    finally:
      builtins.__import__ = real_import
    body = docx_text(p)
    self.assertIn("TAILMARK", body)
    self.assertIn("ALTMARK", body)
    self.assertTrue(any("SVG" in str(w.message) for w in caught))

  def test_missing_rasterizer_keeps_the_logo_render_alive(self):
    import builtins
    real_import = builtins.__import__

    def blocking_import(name, *a, **kw):
      if name in ("rlPyCairo", "svglib", "svglib.svglib"):
        raise ImportError(f"No module named {name!r}")
      return real_import(name, *a, **kw)

    p = self.path()
    svg = self.write_svg()
    builtins.__import__ = blocking_import
    try:
      with warnings.catch_warnings(record=True):
        warnings.simplefilter("always")
        md_to_docx(f"---\ntitle: T\nlogo: {self.md_url(svg)}\n---\n\nBODYMARK.\n", p)
    finally:
      builtins.__import__ = real_import
    self.assertIn("BODYMARK", docx_text(p))

  def test_unreadable_svg_degrades_to_alt_text(self):
    """A broken figure costs its own figure, never the whole render."""
    p = self.path()
    broken = os.path.join(self._tmp.name, "broken.svg")
    with open(broken, "w", encoding="utf-8") as f:
      f.write("this is not svg at all")
    md_to_docx(f"![ALTMARK]({self.md_url(broken)})\n\nTAILMARK.\n", p)
    body = docx_text(p)
    self.assertIn("TAILMARK", body)
    self.assertIn("ALTMARK", body)

#----------------------------------------------------------------------------------- List numbering

class TestOrderedListNumbering(TempDocxCase):
  """`numId` is the counter Word increments, so each ordered list needs its own
  instance to start over, and a CommonMark `start` has to reach the file."""

  def test_second_list_restarts(self):
    p = self.path()
    md_to_docx("1. alpha\n2. beta\n\nbetween\n\n1. gamma\n2. delta\n", p)
    ids = {text: num for text, _, num in docx_numbered_paras(p)}
    self.assertIsNotNone(ids["alpha"])
    self.assertEqual(ids["alpha"], ids["beta"])
    self.assertEqual(ids["gamma"], ids["delta"])
    self.assertNotEqual(ids["alpha"], ids["gamma"])

  def test_commonmark_start_is_honored(self):
    p = self.path()
    md_to_docx("5. five\n6. six\n", p)
    self.assertIn("5", docx_start_overrides(p).values())

  def test_nested_list_leaves_the_outer_counter_alone(self):
    p = self.path()
    md_to_docx("1. outer one\n2. outer two\n   1. inner one\n3. outer three\n", p)
    ids = {text: num for text, _, num in docx_numbered_paras(p)}
    self.assertEqual(ids["outer one"], ids["outer three"])
    self.assertNotEqual(ids["outer one"], ids["inner one"])

  def test_bullets_keep_style_numbering(self):
    p = self.path()
    md_to_docx("- a\n- b\n", p)
    for text, style, num_id in docx_numbered_paras(p):
      self.assertTrue(style.startswith("List Bullet"))
      self.assertIsNone(num_id)

  def test_deep_level_does_not_duplicate_the_item(self):
    """A style the document lacks must not leave a stray paragraph behind."""
    p = self.path()
    d = DOCX(p)
    d.bullet("LVLTEXT", level=3)
    d.save()
    texts = [t for t, _, _ in docx_numbered_paras(p)]
    self.assertEqual(1, texts.count("LVLTEXT"))

#--------------------------------------------------------------------------------------------- Runs

class TestRunFormatting(TempDocxCase):
  """`text()` documents `link_url` among its kwargs, and a run appended to a
  heading belongs to that heading's type."""

  def test_link_url_produces_a_hyperlink(self):
    p = self.path()
    d = DOCX(p)
    d.text("clickme", link_url="https://example.com")
    d.save()
    self.assertIn("hyperlink", docx.Document(p).element.xml)

  def test_run_after_heading_inherits_heading_font(self):
    p = self.path()
    d = DOCX(p)
    d.heading("HEAD", 1).text(" appended")
    d.save()
    for para in docx.Document(p).paragraphs:
      if "HEAD" not in para.text:
        continue
      for run in para.runs:
        self.assertIsNone(run.font.size, "appended run pins a body size")
        self.assertIsNone(run.font.name, "appended run pins a body family")

#-------------------------------------------------------------------------------------- OOXML order

class TestParagraphPropertyOrder(TempDocxCase):
  """`pPr` children are a sequence in CT_PPrBase: `pBdr` and `shd` come before
  `spacing` and `ind`, or a strict renderer may drop them."""

  ORDER = ("pStyle", "numPr", "pBdr", "shd", "spacing", "ind", "contextualSpacing")

  def test_border_and_shading_precede_spacing(self):
    p = self.path()
    md_to_docx("```py\nx = 1\n```\n\n> quoted text\n", p)
    groups = docx_ppr_children(p)
    self.assertTrue(groups, "no bordered or shaded paragraph produced")
    for kids in groups:
      ranked = [self.ORDER.index(k) for k in kids if k in self.ORDER]
      self.assertEqual(sorted(ranked), ranked, f"out of schema order: {kids}")

#-------------------------------------------------------------------------------------- Frontmatter

class TestFrontmatter(TempDocxCase):
  """A leading `---` is a thematic break as often as a frontmatter opener. Only
  a YAML mapping is frontmatter; the rest is content and stays."""

  def test_leading_rule_keeps_its_content(self):
    p = self.path()
    md_to_docx("---\n# Intro\nBODYTEXT here\n---\n# Section 2\n", p)
    body = docx_text(p)
    for marker in ("Intro", "BODYTEXT", "Section 2"):
      self.assertIn(marker, body)

  def test_real_frontmatter_is_still_consumed(self):
    p = self.path()
    md_to_docx("---\ntitle: Doc Title\nauthor: sara\n---\n\nBODYTEXT.\n", p)
    body = docx_text(p)
    self.assertIn("BODYTEXT", body)
    self.assertNotIn("author: sara", body)

  def test_broken_yaml_warns(self):
    with warnings.catch_warnings(record=True) as caught:
      warnings.simplefilter("always")
      md_to_docx("---\nauthor: [unclosed\ntitle: X\n---\n\nBody.\n", self.path())
    self.assertTrue(any("frontmatter" in str(w.message) for w in caught))

#------------------------------------------------------------------------------------------ Mermaid

class TestMermaidNetwork(unittest.TestCase):
  """The mermaid.ink fallback sends the diagram source to a third party, so it
  answers to a switch."""

  def test_remote_false_makes_no_request(self):
    import docmarq.md.mermaid as mermaid
    attempts = []
    real = mermaid._try_mermaid_ink
    mermaid._try_mermaid_ink = lambda *a, **kw: (attempts.append(1), False)[1]
    try:
      mermaid.compile_to_png("flowchart LR\n  X-->Y", cli="no-such-mmdc-binary",
        scale=3, remote=False)
      self.assertEqual(0, len(attempts))
      mermaid.compile_to_png("flowchart LR\n  X-->Y", cli="no-such-mmdc-binary",
        scale=3, remote=True)
      self.assertEqual(1, len(attempts))
    finally:
      mermaid._try_mermaid_ink = real

  def test_truncated_cache_entry_is_rejected(self):
    import docmarq.md.mermaid as mermaid
    path = mermaid._cache_path("testtesttesttest")
    try:
      path.write_bytes(b"\x89PNG\r\n\x1a\n truncated")
      self.assertFalse(mermaid._valid_png(path))
    finally:
      path.unlink(missing_ok=True)

  def test_cache_key_matches_pdfmarq(self):
    """The README promises a cache shared with pdfmarq; the key is that promise."""
    try:
      import pdfmarq.md.mermaid as other
    except ImportError:
      self.skipTest("pdfmarq not installed")
    import docmarq.md.mermaid as mine
    args = ("flowchart LR\nA-->B", "default", "transparent", 3,
      "IBMPlexSans", "./fonts", "mmdc")
    self.assertEqual(other._cache_key(*args), mine._cache_key(*args))
    self.assertEqual(str(other._CACHE_DIR), str(mine._CACHE_DIR))

if __name__ == "__main__":
  unittest.main(verbosity=2)
