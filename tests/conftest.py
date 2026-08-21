# tests/conftest.py

"""With `python_functions = ["*"]`, collect only functions defined in the test module; plus the shared DOCX validator."""

import inspect
from pathlib import Path

def pytest_pycollect_makeitem(collector, name, obj):
  if inspect.isfunction(obj) and obj.__module__ != collector.obj.__name__:
    return [] # ignore library functions imported into the test file

#--------------------------------------------------------------------------------------- Assertions

def assert_valid_docx(path:str|Path, min_size:int=2000):
  """Validate that `path` points to a real-looking DOCX file (ZIP magic `PK\\x03\\x04`)."""
  p = Path(path)
  assert p.exists(), f"DOCX not created: {p}"
  size = p.stat().st_size
  assert size >= min_size, f"DOCX suspiciously small ({size} bytes): {p}"
  with p.open("rb") as f:
    head = f.read(4)
  assert head == b"PK\x03\x04", f"Not a DOCX/ZIP (header={head!r}): {p}"

#--------------------------------------------------------------------------------------- Inspection

def docx_text(path:str) -> str:
  """Every string in the document, table cells included."""
  import docx
  d = docx.Document(path)
  parts = [p.text for p in d.paragraphs]
  for table in d.tables:
    for row in table.rows:
      parts += [c.text for c in row.cells]
  return "\n".join(parts)

def docx_numbered_paras(path:str) -> list:
  """`(text, style, numId)` for every non-empty paragraph."""
  import docx
  from docx.oxml.ns import qn
  out = []
  for p in docx.Document(path).paragraphs:
    if not p.text:
      continue
    num_id = None
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
      numPr = pPr.find(qn("w:numPr"))
      if numPr is not None:
        el = numPr.find(qn("w:numId"))
        num_id = el.get(qn("w:val")) if el is not None else None
    out.append((p.text, p.style.name, num_id))
  return out

def docx_start_overrides(path:str) -> dict:
  """`numId → startOverride` for every numbering instance that declares one."""
  import docx
  from docx.oxml.ns import qn
  found = {}
  numbering = docx.Document(path).part.numbering_part.element
  for num in numbering.findall(qn("w:num")):
    override = num.find(qn("w:lvlOverride"))
    if override is None:
      continue
    start = override.find(qn("w:startOverride"))
    if start is not None:
      found[num.get(qn("w:numId"))] = start.get(qn("w:val"))
  return found

def docx_ppr_children(path:str) -> list:
  """Child-tag lists of every `pPr` that carries a border or shading."""
  import docx
  from docx.oxml.ns import qn
  out = []
  for p in docx.Document(path).paragraphs:
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
      continue
    kids = [c.tag.split("}")[1] for c in pPr]
    if "pBdr" in kids or "shd" in kids:
      out.append(kids)
  return out
