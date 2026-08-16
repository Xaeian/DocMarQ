# docmarq/md/renderer.py

"""Markdown → DOCX rendering.

Walks `markdown-it-py` tokens and emits `DOCX` API calls. Word's native
layout (auto page breaks, keep-with-next on Heading styles, paragraph
spacing) does the heavy lifting - this renderer is a token dispatcher
with an inline state machine.

Pure helpers live in sibling modules: `tokens.py` (attr/find_close),
`slug.py` (heading slugs and anchor resolution), `image_utils.py` (PIL
preprocess + scaling), `mermaid.py` (mmdc CLI). Class methods below own
everything that touches `doc` state.
"""
import os
from markdown_it import MarkdownIt
from markdown_it.token import Token
from mdit_py_plugins.footnote import footnote_plugin
from docx.shared import Pt, RGBColor
from ..core import DOCX
from ..constants import Align
from .style import MarkdownStyle
from .tokens import get_attr, find_close, CALLOUT_RE
from . import slug, image_utils, mermaid

#--------------------------------------------------------------------------- Frontmatter helpers

def strip_frontmatter(text:str) -> tuple[dict|None, str]:
  """Detect and strip YAML frontmatter at the start of the document.

  Returns `(parsed_dict, remaining_text)`. When PyYAML isn't installed
  the block is still stripped but `parsed_dict` is `None` (metadata
  won't auto-fill but the content renders cleanly).
  """
  if not text.startswith("---"):
    return None, text
  parts = text.split("\n", 1)
  if parts[0].strip() != "---" or len(parts) < 2:
    return None, text
  rest = parts[1]
  # Find closing `---` on its own line. Tolerate `\r\n` line endings.
  end_match = None
  for sep in ("\n---\n", "\n---\r\n", "\n---"):
    idx = rest.find(sep)
    if idx != -1:
      end_match = (idx, idx + len(sep))
      break
  if end_match is None:
    return None, text
  yaml_block = rest[:end_match[0]]
  remainder = rest[end_match[1]:].lstrip("\n").lstrip("\r\n")
  try:
    import yaml
  except ImportError:
    import warnings
    warnings.warn(
      "YAML frontmatter needs PyYAML (pip install docmarq[md]); "
      "title/banner/metadata are skipped",
      RuntimeWarning, stacklevel=2,
    )
    return None, remainder
  try:
    data = yaml.safe_load(yaml_block) or {}
  except Exception:
    data = None
  return data, remainder

def peek_frontmatter(text:str) -> dict|None:
  """Return parsed frontmatter without consuming it. Useful for early
  metadata inspection (e.g. `landscape:` flag) before constructing the
  document. Symmetric with `pdfmarq.md.peek_frontmatter`."""
  fm, _ = strip_frontmatter(text)
  return fm

#----------------------------------------------------------------------------- MarkdownRenderer

class MarkdownRenderer:
  """Render a `markdown-it` token stream onto a `DOCX` instance.

  Stateful (`_list_depth` tracks nesting), so reuse only within a single
  `render()` pass.
  """
  def __init__(self, doc:DOCX, style:MarkdownStyle|None=None,
      base_dir:str|None=None, font_dir:str|None=None):
    """Create a renderer bound to `doc`.

    Args:
      doc: Target `DOCX` instance.
      style: Optional `MarkdownStyle`. `None` uses defaults.
      base_dir: Root for resolving relative image paths
        (e.g. `![alt](./img/x.png)`). Defaults to current working dir.
      font_dir: Optional TTF root. When set, mermaid diagrams render with
        `style.body_family` instead of the system default sans-serif
        (matches the rest of the document). Layout: `<font_dir>/<family>/
        <family>-Regular.ttf` (mirrors `pdfmarq.FontManager`).
    """
    self.doc = doc
    self.style = style or MarkdownStyle()
    self.base_dir = base_dir or os.getcwd()
    self.font_dir = font_dir
    md = MarkdownIt("commonmark", {"html": True, "breaks": False})
    md.enable(["table", "strikethrough"])
    md.use(footnote_plugin)
    # `$...$` / `$$...$$` → `math_inline` / `math_block` tokens (OMML).
    # Degrades gracefully when missing - `$...$` stays as literal text.
    self._math_enabled = False
    if self.style.math_enable:
      try:
        from mdit_py_plugins.dollarmath import dollarmath_plugin
        md.use(dollarmath_plugin)
        self._math_enabled = True
      except ImportError:
        import warnings
        warnings.warn(
          "math formulas need `mdit-py-plugins` (pip install docmarq[md]); "
          "`$...$` will render as plain text",
          RuntimeWarning, stacklevel=2,
        )
    self._md = md
    self._list_depth = 0
    # Pre-scanned slugs for `[text](#slug)` internal links - built before
    # the first token so broken anchors render as plain text, never as
    # dangling Word hyperlinks.
    self._known_slugs: set[str] = set()

  #-------------------------------------------------------------------------------- Entry point

  def render(self, md_text:str):
    """Parse markdown and emit it as DOCX content. Strips YAML frontmatter
    if present; recognized metadata keys auto-fill `doc.metadata()`. When
    frontmatter is present and `style.banner_render` is on, renders a
    first-page banner (title, status, author, dates) before body content.
    """
    fm, md_text = self._strip_frontmatter(md_text)
    if fm:
      self._apply_frontmatter_metadata(fm)
    self._apply_chrome(fm)
    # Body typography applied once; caller may override via `doc.style(...)`.
    self.doc.style(
      line_height=self.style.line_height,
      space_after=self.style.para_gap_pt,
      code_family=self.style.mono_family,
      code_color=self.style.code_inline_color,
      code_bg=self.style.code_inline_bg,
    )
    if fm and self.style.banner_render:
      self._render_banner(fm)
    tokens = self._md.parse(md_text)
    # Pre-scan so `link_open` can validate `#anchor` hrefs before rendering.
    self._known_slugs = slug.collect_heading_slugs(tokens)
    self._render_tokens(tokens)

  #----------------------------------------------------------------------------- Page chrome

  def _apply_chrome(self, fm:dict|None):
    """Set up footer (page numbers, every page) and continuation header
    (short title/id on page 2+, page 1 stays clean for the banner).

    Word's `titlePg` flag is section-wide - it splits BOTH header and
    footer between first-page and default variants. To keep page numbers
    on every page we have to set first-page footer AND default footer.
    """
    s = self.style
    sec = self.doc._doc.sections[-1]
    footer_template = None
    if s.page_number_label:
      footer_template = f"{s.page_number_label} {{page}}"
      if s.page_number_total:
        footer_template += " / {pages}"
    head_text = None  # id wins over title for the continuation header
    if s.mini_banner_render and fm:
      head_text = fm.get("id") or fm.get("title")
      if head_text:
        head_text = str(head_text)
    # Split first/default only when the banner occupies page 1 - keeps
    # the continuation header off the banner page.
    use_different_first = bool(fm and s.banner_render and head_text)
    if use_different_first:
      sec.different_first_page_header_footer = True
      sec.first_page_header.paragraphs[0].text = ""  # banner is the masthead
      self._fill_header(sec.header.paragraphs[0], head_text)
      if footer_template:
        self._fill_footer(sec.first_page_footer.paragraphs[0], footer_template)
        self._fill_footer(sec.footer.paragraphs[0], footer_template)
    else:
      if head_text:
        self._fill_header(sec.header.paragraphs[0], head_text)
      if footer_template:
        self._fill_footer(sec.footer.paragraphs[0], footer_template)

  @staticmethod
  def _fill_header(p, text:str):
    """Set a header paragraph to `text` - centered, muted grey, 9pt."""
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x6E, 0x77)

  @staticmethod
  def _fill_footer(p, template:str):
    """Set a footer paragraph using the `{page}`/`{pages}` template form,
    centered. Delegates to the docx core's field-injection helper."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from ..core import _inject_page_fields
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _inject_page_fields(p, template)

  #----------------------------------------------------------------------- First-page banner

  def _render_banner(self, fm:dict):
    """Draw a paper-style first-page masthead from YAML frontmatter.

    Two layouts:
      - With `logo: file.png` in frontmatter → 1×2 borderless table, logo
        on the left, title/status/author/dates on the right.
      - Without logo → single-column flow _(simpler, matches pdfmarq when
        no logo is supplied)_.

    The horizontal rule separating banner from body is always emitted.
    """
    logo_path = self._resolve_logo_path(fm.get("logo"))
    if logo_path:
      self._render_banner_table(fm, logo_path)
    else:
      self._render_banner_flow(fm)
    self.doc.hr(color=self.style.hr_color)

  def _resolve_logo_path(self, logo) -> str|None:
    """Resolve `logo:` to a local file path. `None` if unset, remote, or
    missing (banner falls back to flow layout). Missing-but-set warns so
    typos don't silently drop the logo. Mirrors `pdfmarq` (twin libs)."""
    if not logo:
      return None
    path = image_utils.resolve_path(str(logo), self.base_dir)
    if path and os.path.isfile(path):
      return path
    import warnings
    warnings.warn(
      f"frontmatter `logo: {logo}` not found "
      f"(base_dir={self.base_dir}); rendering without logo",
      RuntimeWarning, stacklevel=2,
    )
    return None

  def _render_banner_flow(self, fm:dict):
    """Single-column banner layout (no logo). Uses the fluent `DOCX` API."""
    from docx.enum.text import WD_LINE_SPACING
    def new_para(*, space_before:float=0, space_after:float=0,
        line_spacing:float|None=None, line_spacing_exact_pt:float|None=None):
      self.doc.para()
      pf = self.doc._current_para.paragraph_format
      pf.space_before = Pt(space_before)
      pf.space_after = Pt(space_after)
      if line_spacing_exact_pt is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing_exact_pt)
      elif line_spacing is not None:
        pf.line_spacing = line_spacing
    def add_run(text, *, bold=False, color=None, bg=None, size=None):
      self.doc.text(text, bold=bold, color=color, bg_color=bg, size=size)
    self._render_banner_content(fm, new_para=new_para, add_run=add_run)

  def _render_banner_table(self, fm:dict, logo_path:str):
    """Two-column banner: logo left, title/meta right. Built as a borderless
    1×2 table - python-docx doesn't support floating images cleanly, so a
    layout table is the simplest path to side-by-side without raw OOXML.
    """
    from ..core import _set_table_widths
    from ..tables import (remove_table_borders, set_cell_margins,
      set_cell_vertical_align)
    self.doc._flush_para()
    content_w = self.doc._page.content_width
    logo_col_w = 30 # mm - leaves room for square-ish logos
    text_col_w = content_w - logo_col_w
    tbl = self.doc._doc.add_table(rows=1, cols=2)
    _set_table_widths(tbl, [logo_col_w, text_col_w])
    remove_table_borders(tbl)
    left_cell, right_cell = tbl.rows[0].cells
    set_cell_margins(left_cell, top=0, right=4, bot=0, left=0)
    set_cell_vertical_align(left_cell, "center")
    self._fill_cell_logo(left_cell, logo_path, max_w_mm=logo_col_w - 4)
    set_cell_margins(right_cell, top=0, right=0, bot=0, left=2)
    set_cell_vertical_align(right_cell, "center")
    self._fill_cell_banner_text(right_cell, fm)

  def _fill_cell_logo(self, cell, logo_path:str, max_w_mm:float):
    """Embed the logo image into a cell, sized to fit `max_w_mm` wide.
    Always preprocesses through Pillow to crop transparent padding - source
    logos commonly have transparent margins that would shift the visible
    content right of the cell edge.
    """
    from docx.shared import Mm
    from ..utils import to_mm
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    from docx.image.exceptions import UnrecognizedImageError
    run = p.add_run()
    w_emu = Mm(to_mm(max_w_mm, self.doc.unit))
    buf = image_utils.preprocess_to_buffer(logo_path)
    try:
      run.add_picture(buf if buf is not None else logo_path, width=w_emu)
      self._pin_inline_picture_offsets(run)
    except (OSError, ValueError, UnrecognizedImageError):
      pass

  def _fill_cell_banner_text(self, cell, fm:dict):
    """Banner-content adapter that writes into a table cell using
    `python-docx` directly. The fluent `doc.text()` API targets the
    document-level current paragraph, so cells need their own callbacks.
    """
    from ..utils import color_hex, rgb255
    from ..inline import _apply_run_shading
    from docx.enum.text import WD_LINE_SPACING
    state = {"used_first": False, "current_p": None}
    def new_para(*, space_before:float=0, space_after:float=0,
        line_spacing:float|None=None, line_spacing_exact_pt:float|None=None):
      if not state["used_first"]:
        p = cell.paragraphs[0]
        p.text = ""
        state["used_first"] = True
      else:
        p = cell.add_paragraph()
      pf = p.paragraph_format
      pf.space_before = Pt(space_before)
      pf.space_after = Pt(space_after)
      if line_spacing_exact_pt is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing_exact_pt)
      elif line_spacing is not None:
        pf.line_spacing = line_spacing
      state["current_p"] = p
    def add_run(text, *, bold=False, color=None, bg=None, size=None):
      run = state["current_p"].add_run(text)
      if bold: run.bold = True
      if size is not None: run.font.size = Pt(size)
      if color is not None:
        run.font.color.rgb = RGBColor(*rgb255(color))
      if bg is not None:
        _apply_run_shading(run, color_hex(bg))
    self._render_banner_content(fm, new_para=new_para, add_run=add_run)

  def _render_banner_content(self, fm:dict, *, new_para, add_run):
    """Banner data rendering. `new_para` / `add_run` callbacks bridge to
    either the fluent doc API (flow layout) or direct cell paragraphs
    (table layout) - the content sequence itself is identical.

    Args:
      fm: Frontmatter dict.
      new_para: `new_para(space_before=0, space_after=0)` opens paragraph.
      add_run: `add_run(text, bold, color, bg, size)` adds styled run.
    """
    s = self.style
    title = fm.get("title")
    if title:
      # `EXACTLY` bypasses Calibri's ~1.2× internal leading so large title
      # text sits tight without gaps between wrapped lines.
      from ..utils import tight_line_height_pt
      new_para(line_spacing_exact_pt=tight_line_height_pt(s.banner_title_size))
      add_run(str(title), bold=True, size=s.banner_title_size)
    status = (fm.get("status") or "").strip().lower()
    version = fm.get("version")
    id_ = fm.get("id")
    if status or version or id_:
      new_para(space_before=2)
      sep = "  •  "  # U+2022 BULLET - more visible than middle-dot
      if status:
        bg, fg = s.banner_status_colors.get(status,
          ((0.9, 0.9, 0.9), (0.3, 0.3, 0.3)))
        add_run(f" {status.upper()} ", bold=True, color=fg, bg=bg)
        if version or id_:
          add_run(sep, color=s.muted_color)
      if version:
        add_run(str(version), color=s.muted_color)
        if id_:
          add_run(sep, color=s.muted_color)
      if id_:
        add_run(str(id_), color=s.muted_color)
    author = fm.get("author")
    if author:
      new_para()
      add_run(f"{s.banner_label_author}: {author}",
        size=s.banner_meta_size, color=s.muted_color)
    created = fm.get("created")
    updated = fm.get("updated")
    if created or updated:
      new_para()
      parts = []
      if created:
        parts.append(f"{s.banner_label_created}: {self._format_date(created)}")
      if updated:
        parts.append(f"{s.banner_label_updated}: {self._format_date(updated)}")
      add_run("  •  ".join(parts),
        size=s.banner_meta_size, color=s.muted_color)

  def _format_date(self, value) -> str:
    """Format a frontmatter date value using `style.date_format`. Accepts
    `datetime.date`/`datetime.datetime` (parsed by PyYAML) or raw strings."""
    if value is None:
      return ""
    if hasattr(value, "strftime"):
      return value.strftime(self.style.date_format)
    return str(value)

  @staticmethod
  def _strip_frontmatter(text:str) -> tuple[dict|None, str]:
    """Detect and strip YAML frontmatter at the start of the document."""
    return strip_frontmatter(text)

  def _apply_frontmatter_metadata(self, fm:dict):
    """Push recognized YAML keys into `doc.metadata()`. Unknown keys silently
    ignored; banner rendering is handled separately."""
    meta = {}
    for yaml_key, meta_key in (("title", "title"), ("author", "author"),
        ("subject", "subject")):
      v = fm.get(yaml_key)
      if v is not None:
        meta[meta_key] = str(v)
    kw = fm.get("keywords")
    if kw is not None:
      if isinstance(kw, (list, tuple)):
        meta["keywords"] = ", ".join(str(k) for k in kw)
      else:
        meta["keywords"] = str(kw)
    if meta:
      self.doc.metadata(**meta)

  #--------------------------------------------------------------------------------- Directives

  def _find_group_close(self, tokens:list[Token], start:int) -> int:
    """Find index of matching `<!-- /group -->` for the open at `start`.
    Tracks depth so nested directives (officially unsupported per #14)
    still let the outer group find its close. Returns `len(tokens)` for
    unclosed groups, warns once."""
    from .tokens import is_group_open_directive, is_group_close_directive
    depth = 1
    for j in range(start + 1, len(tokens)):
      t = tokens[j]
      if t.type != "html_block": continue
      content = t.content or ""
      if is_group_open_directive(content): depth += 1
      elif is_group_close_directive(content):
        depth -= 1
        if depth == 0: return j
    import warnings
    warnings.warn(
      "unclosed `<!-- group -->` directive - rendering to end of document",
      RuntimeWarning, stacklevel=2,
    )
    return len(tokens)

  def _render_group(self, tokens:list[Token]):
    """Render group content with Word `keep_with_next` chaining.

    Word handles pagination itself - we tell it to keep our paragraphs
    together via `<w:keepNext/>` on each paragraph except the last in
    the group. Tables embedded in the group don't get special handling
    (no clean public API for whole-table keep-together); their leading
    paragraph's `keep_with_next` keeps them tethered to it, which covers
    the common case well enough for MVP.
    """
    if not tokens: return
    self.doc._flush_para()
    start_idx = len(self.doc._doc.paragraphs)
    self._render_tokens(tokens)
    self.doc._flush_para()
    added = self.doc._doc.paragraphs[start_idx:]
    # Set keep_with_next on every paragraph except the last - that chains
    # them so Word's paginator treats the run as one logical block.
    for p in added[:-1]:
      p.paragraph_format.keep_with_next = True

  #---------------------------------------------------------------------------- Block dispatch

  def _render_tokens(self, tokens:list[Token]):
    i = 0
    n = len(tokens)
    while i < n:
      t = tokens[i]
      tt = t.type
      if tt == "heading_open":
        level = int(t.tag[1])
        self._render_heading(tokens[i+1], level)
        i += 3
      elif tt == "paragraph_open":
        # Standalone images get block treatment (embedded + centered);
        # mixed inline content uses regular paragraph rendering.
        inline = tokens[i+1]
        img_src = self._standalone_image_src(inline)
        if img_src:
          self._render_block_image(img_src, inline)
        else:
          self._render_paragraph(inline)
        i += 3
      elif tt == "fence" or tt == "code_block":
        info_parts = (t.info or "").strip().split(maxsplit=1)  # [lang, rest_dsl]
        lang = info_parts[0] if info_parts else None
        info_rest = info_parts[1] if len(info_parts) > 1 else ""
        if lang == "mermaid": self._render_mermaid(t.content, info_rest)
        elif lang == "math": self._render_math_block(t.content)
        else:
          self.doc.code_block(
            t.content,
            language=lang,
            bg_color=self.style.code_block_bg,
            border_color=self.style.code_block_border,
            font_family=self.style.mono_family,
          )
        i += 1
      elif tt == "math_block":
        self._render_math_block(t.content)
        i += 1
      elif tt == "bullet_list_open": i = self._render_list(tokens, i, ordered=False)
      elif tt == "ordered_list_open": i = self._render_list(tokens, i, ordered=True)
      elif tt == "blockquote_open": i = self._render_blockquote(tokens, i)
      elif tt == "hr":
        self.doc.hr(color=self.style.hr_color)
        i += 1
      elif tt == "table_open": i = self._render_table(tokens, i)
      elif tt == "footnote_block_open": i = self._render_footnote_block(tokens, i)
      elif tt == "html_block":
        # Only directive comments are handled; arbitrary HTML is dropped.
        from .tokens import (is_pagebreak_directive,
          is_group_open_directive, is_group_close_directive)
        content = t.content or ""
        if is_pagebreak_directive(content):
          self.doc.page_break()
        elif is_group_open_directive(content):
          end_i = self._find_group_close(tokens, i)
          self._render_group(tokens[i+1:end_i])
          i = end_i + 1
          continue
        elif is_group_close_directive(content):
          import warnings
          warnings.warn(
            "stray `<!-- /group -->` directive (no matching open)",
            RuntimeWarning, stacklevel=2,
          )
        i += 1
      else:
        i += 1

  #-------------------------------------------------------------------------------- Headings

  def _render_heading(self, inline_token:Token, level:int):
    """Emit a heading. Bold/italic/strike pass through; font family/size/color
    inherit from the Heading paragraph style (so inline `code` keeps heading
    size, just switches family). Registers the slug as a Word bookmark so
    `[text](#slug)` internal links resolve here.
    """
    self.doc.heading(level=level)
    p = self.doc._current_para
    anchor = slug.slugify_inline(inline_token)
    if anchor:
      self.doc.bookmark(anchor)
    self._add_runs_to_heading(p, inline_token)

  def _add_runs_to_heading(self, p, inline_token:Token):
    """Walk inline children and add runs to a Heading paragraph. Inherits
    size/color from the Heading style; overrides font when `head_family` set."""
    state = {"bold": False, "italic": False, "strike": False}
    head_family = self.style.head_family
    for c in (inline_token.children or []):
      ct = c.type
      if ct == "text":
        run = p.add_run(c.content)
        if head_family: run.font.name = head_family
        if state["bold"]: run.bold = True
        if state["italic"]: run.italic = True
        if state["strike"]: run.font.strike = True
      elif ct == "code_inline":
        run = p.add_run(c.content)
        run.font.name = self.style.mono_family
        if state["bold"]: run.bold = True
        if state["italic"]: run.italic = True
        if state["strike"]: run.font.strike = True
      elif ct == "strong_open": state["bold"] = True
      elif ct == "strong_close": state["bold"] = False
      elif ct == "em_open": state["italic"] = True
      elif ct == "em_close": state["italic"] = False
      elif ct == "s_open": state["strike"] = True
      elif ct == "s_close": state["strike"] = False
      elif ct == "math_inline":
        # Heading math inherits the heading run size (size_halfpt=None) so it
        # scales with the heading rather than rendering at body size.
        _, color_hex = self._math_run_props()
        self._append_inline_math(p, c.content, None, color_hex)
      elif ct == "softbreak": p.add_run(" ")
      elif ct == "hardbreak": p.add_run("\n")

  #-------------------------------------------------------------------------------- Paragraph

  def _render_paragraph(self, inline_token:Token):
    """Emit a regular body paragraph and stream styled runs into it."""
    self.doc.para()
    self._add_runs_to_current_para(inline_token)

  def _add_runs_to_current_para(self, inline_token:Token):
    """State-machine walk of inline children. Translates each child into
    one or more `doc.text()` / `doc.link()` / `doc.line_break()` calls.
    """
    state = {"bold": False, "italic": False, "strike": False,
      "link_url": None, "link_target": None}

    def kwargs():
      k = {}
      if state["bold"]: k["bold"] = True
      if state["italic"]: k["italic"] = True
      if state["strike"]: k["strike"] = True
      return k

    for c in (inline_token.children or []):
      ct = c.type
      if ct == "text":
        if state["link_target"]:
          self.doc.link(c.content, target=state["link_target"])
        elif state["link_url"]:
          self.doc.link(c.content, url=state["link_url"])
        else:
          self.doc.text(c.content, **kwargs())
      elif ct == "code_inline":
        # Inherits wrapping bold/italic/strike - e.g. `*foo `code` bar*`.
        self.doc.text(c.content, code=True, **kwargs())
      elif ct == "strong_open": state["bold"] = True
      elif ct == "strong_close": state["bold"] = False
      elif ct == "em_open": state["italic"] = True
      elif ct == "em_close": state["italic"] = False
      elif ct == "s_open": state["strike"] = True
      elif ct == "s_close": state["strike"] = False
      elif ct == "link_open":
        href = get_attr(c, "href") or ""
        # `#anchor` validated against pre-scanned slugs; unknown → plain text.
        anchor = slug.resolve_anchor(href, self._known_slugs)
        if anchor is not None:
          state["link_target"] = anchor
        else:
          state["link_url"] = self._resolve_link(href)
      elif ct == "link_close":
        state["link_url"] = None
        state["link_target"] = None
      elif ct == "math_inline":
        self._render_inline_math(c.content)
      elif ct == "softbreak": self.doc.text(" ")
      elif ct == "hardbreak": self.doc.line_break()
      elif ct == "image":
        # Inline images render as italic alt text; full embedding is deferred.
        alt = c.content or "[image]"
        self.doc.text(alt, italic=True)
      elif ct == "footnote_ref":
        # `[^N]` → superscript `[1]` link style.
        label = (c.meta or {}).get("label") or ""
        self.doc.text(f"[{label}]", superscript=True, color=self.style.link_color)

  def _resolve_link(self, href:str) -> str|None:
    """Resolve local hrefs against `link_root`/`link_base`. URLs with a
    schema or hash-anchor are returned as-is. `None` means "render text
    but don't linkify" (for unresolvable local refs).
    """
    if not href:
      return None
    if href.startswith(("http://", "https://", "mailto:", "ftp://")):
      return href
    if href.startswith("#"):
      return None  # internal anchors handled via bookmark targets
    if self.style.link_root:
      base = self.style.link_base.strip("/")
      root = self.style.link_root.rstrip("/")
      if href.startswith("/"):
        return f"{root}{href}"
      return f"{root}/{base}/{href}" if base else f"{root}/{href}"
    return None

  #------------------------------------------------------------------------------------ Lists

  def _render_list(self, tokens:list[Token], start:int, ordered:bool) -> int:
    open_type = "ordered_list_open" if ordered else "bullet_list_open"
    close_type = "ordered_list_close" if ordered else "bullet_list_close"
    end = find_close(tokens, start, open_type, close_type)
    depth = self._list_depth
    self._list_depth += 1
    j = start + 1
    while j < end:
      if tokens[j].type == "list_item_open":
        item_end = find_close(tokens, j, "list_item_open", "list_item_close")
        self._render_list_item(tokens, j + 1, item_end, ordered, depth)
        j = item_end + 1
      else:
        j += 1
    self._list_depth -= 1
    return end + 1

  def _render_list_item(self, tokens:list[Token], start:int, end:int,
      ordered:bool, depth:int):
    """Render one list item. First inline content goes into a list-styled
    paragraph; nested lists recurse with deeper depth."""
    first_block = True
    k = start
    while k < end:
      t = tokens[k]
      tt = t.type
      if tt == "paragraph_open":
        inline = tokens[k+1]
        if first_block:
          if ordered:
            self.doc.ordered(level=depth)
          else:
            self.doc.bullet(level=depth)
          self._add_runs_to_current_para(inline)
          first_block = False
        else:
          # Continuation paragraph inside a list item - plain para
          self._render_paragraph(inline)
        k += 3
      elif tt == "bullet_list_open": k = self._render_list(tokens, k, ordered=False)
      elif tt == "ordered_list_open": k = self._render_list(tokens, k, ordered=True)
      else:
        k += 1

  #-------------------------------------------------------------------------------- Blockquote

  @staticmethod
  def _spacing_for(idx:int, total:int, before:float, after:float) -> tuple[float, float]:
    """Pick `(space_before, space_after)` for paragraph `idx` in a group of
    `total`. First gets `before`, last gets `after`, inner gets 0/0 - so
    the left bar reads as one continuous block rather than stacked stripes."""
    if total <= 1:
      return (before, after)
    if idx == 0:
      return (before, 0)
    if idx == total - 1:
      return (0, after)
    return (0, 0)

  # Quote content is flattened to `inline` tokens; a fence carries its text on
  # `token.content` instead, so it would vanish. Reported, not dropped silently.
  _QUOTE_UNSUPPORTED = ("fence", "code_block")

  def _collect_blockquote_inlines(self, tokens:list[Token], start:int, end:int) -> list[Token]:
    """Flat list of every `inline` token directly inside the blockquote.
    Nested blockquotes / lists / code blocks not yet supported - they'd
    require recursion and per-block grouping."""
    out = []
    dropped = []
    j = start + 1
    while j < end:
      t = tokens[j]
      if t.type == "inline": out.append(t)
      elif t.type in self._QUOTE_UNSUPPORTED: dropped.append(t)
      j += 1
    if dropped:
      import warnings
      warnings.warn(
        f"{len(dropped)} code block(s) inside a blockquote are not rendered; "
        "move them outside the `>` quote to keep the content",
        RuntimeWarning, stacklevel=2,
      )
    return out

  def _render_blockquote(self, tokens:list[Token], start:int) -> int:
    end = find_close(tokens, start, "blockquote_open", "blockquote_close")
    callout_type = self._detect_callout(tokens, start, end)  # `> [!NOTE]` form
    if callout_type:
      return self._render_callout(tokens, start, end, callout_type)
    inlines = self._collect_blockquote_inlines(tokens, start, end)
    total = len(inlines)
    for idx, inline in enumerate(inlines):
      sb, sa = self._spacing_for(idx, total, 3, 3)
      self.doc.blockquote(text=None,
        border_color=self.style.quote_border,
        text_color=self.style.muted_color,
        space_before=sb, space_after=sa)
      self._add_runs_to_current_para(inline)
    return end + 1

  def _detect_callout(self, tokens:list[Token], start:int, end:int) -> str|None:
    """Return callout type ('note', 'tip', ...) if the blockquote opens
    with `[!TYPE]`, else `None`. Looks at the first inline token only."""
    j = start + 1
    while j < end:
      if tokens[j].type == "inline":
        content = tokens[j].content or ""
        first_line = content.split("\n", 1)[0]
        m = CALLOUT_RE.match(first_line)
        if m:
          return m.group(1).lower()
        return None
      j += 1
    return None

  def _render_callout(self, tokens:list[Token], start:int, end:int,
      kind:str) -> int:
    """Render a GitHub callout: colored title row + body in muted blockquote.
    All paragraphs in the group share tight inner spacing so the left bar
    reads as a single continuous block regardless of body length.
    """
    label = getattr(self.style, f"callout_label_{kind}", kind.title())
    border_rgb, text_rgb = self.style.callout_colors.get(kind,
      (self.style.quote_border, self.style.muted_color))
    inlines = self._collect_blockquote_inlines(tokens, start, end)
    # Strip the `[!TYPE]` marker from the FIRST inline's children. markdown-it
    # puts marker + body into a single inline (soft-broken), so we can't just
    # skip the whole inline - we'd lose the body that follows the marker line.
    body_inlines = []
    if inlines:
      stripped = self._strip_callout_marker(inlines[0])
      if stripped is not None:
        body_inlines.append(stripped)
      body_inlines.extend(inlines[1:])
    total = 1 + len(body_inlines)
    sb, sa = self._spacing_for(0, total, 3, 3)
    self.doc.blockquote(text=None, border_color=border_rgb, text_color=text_rgb,
      space_before=sb, space_after=sa)
    self.doc.text(label, bold=True, color=text_rgb)
    for j, inline in enumerate(body_inlines):
      sb, sa = self._spacing_for(j + 1, total, 3, 3)
      self.doc.blockquote(text=None, border_color=border_rgb,
        text_color=self.style.muted_color,
        space_before=sb, space_after=sa)
      self._add_runs_to_current_para(inline)
    return end + 1

  @staticmethod
  def _strip_callout_marker(inline_token:Token) -> Token|None:
    """Return a shallow copy of `inline_token` with the leading `[!TYPE]`
    marker text + the softbreak right after it removed from `children`.
    Returns `None` if no real content remains after stripping.
    """
    children = inline_token.children or []
    if not children:
      return None
    new_children = list(children)
    # Drop the `[!TYPE]` text token if present; leave unchanged if it doesn't match.
    if new_children[0].type == "text" and CALLOUT_RE.match(new_children[0].content):
      new_children.pop(0)
      # Drop the following softbreak so the body doesn't start with a blank line.
      if new_children and new_children[0].type in ("softbreak", "hardbreak"):
        new_children.pop(0)
    if not new_children:
      return None
    # `Token(type, tag, nesting)` - copy remaining fields manually.
    clone = type(inline_token)(inline_token.type, inline_token.tag,
      inline_token.nesting)
    clone.children = new_children
    clone.content = inline_token.content
    clone.attrs = inline_token.attrs
    return clone

  #-------------------------------------------------------------------------------------- Math

  def _math_run_props(self) -> tuple[int, str]:
    """`(size_halfpt, color_hex)` for OMML runs matching body font size/color."""
    from ..constants import Defaults
    from ..utils import color_hex
    body_pt = self.doc._style.font_size or Defaults.FONT_SIZE
    return round(body_pt * 2), color_hex(self.style.body_color)

  def _render_inline_math(self, latex:str):
    """Inline `$...$`: append a native OMML equation to the current paragraph,
    falling back to an inline image when the LaTeX is outside the OMML subset.
    """
    if self.doc._current_para is None:
      self.doc.para()
    size_halfpt, color_hex = self._math_run_props()
    self._append_inline_math(self.doc._current_para, latex, size_halfpt, color_hex)

  def _append_inline_math(self, p, latex:str, size_halfpt:int|None, color_hex:str):
    """Append inline math to paragraph `p` (works for body paragraphs and
    headings). Native OMML first; image fallback on conversion failure.

    `size_halfpt=None` lets the math runs inherit the paragraph's run size -
    used for headings so `# Result $x^2$` sizes the formula to the heading,
    not the body. Body callers pass the explicit body size so a non-default
    `font_size` doesn't fall back to Word's 11pt Normal default.
    """
    from .math import latex_to_omath, MathConversionError
    try:
      omath = latex_to_omath(latex, size_halfpt=size_halfpt,
        color_hex=color_hex, display=False)
      p._p.append(omath)
      return
    except (MathConversionError, RecursionError):
      # RecursionError can escape the depth guard through nested sub-parsers.
      pass
    self._append_inline_math_image(p, latex)

  def _render_math_block(self, latex:str):
    """Display `$$...$$` (or ```math): a centered native OMML equation on its
    own paragraph, with an image fallback for unsupported LaTeX."""
    from .math import build_omath_para, MathConversionError
    from ..constants import Align
    from ..utils import align_to_docx
    self.doc._flush_para()
    size_halfpt, color_hex = self._math_run_props()
    try:
      omath_para = build_omath_para(latex, size_halfpt=size_halfpt,
        color_hex=color_hex, align="center")
    except (MathConversionError, RecursionError):
      self._render_math_block_image(latex)
      return
    p = self.doc._doc.add_paragraph()
    self.doc._apply_para_spacing(p)
    p.alignment = align_to_docx(Align.CENTER)
    p._p.append(omath_para)
    self.doc._track_block_spacing(self.doc._style.space_after or 0)

  #------------------------------------------------------------------------- Math image fallback

  def _append_inline_math_image(self, p, latex:str):
    """Render `latex` to a PNG and embed it inline in paragraph `p`, dropped
    onto the text baseline. Last resort (no matplotlib, or mathtext also can't
    parse it): render the LaTeX body as inline `code` - never the raw `$...$`
    delimiters, so the no-`$`-leak contract holds."""
    from .math import render_math_png
    from ..constants import Defaults
    body_pt = self.doc._style.font_size or Defaults.FONT_SIZE
    result = render_math_png(latex, fontsize_pt=body_pt,
      color=self.style.body_color, fontset=self.style.math_fontset)
    if result is None:
      self._warn_math_fallback()
      self.doc.text(latex, code=True)
      return
    buf, _w_mm, h_mm, baseline_mm = result
    from docx.shared import Mm
    from ..utils import to_mm
    run = p.add_run()
    try:
      run.add_picture(buf, height=Mm(to_mm(h_mm, self.doc.unit)))
    except Exception:
      self._warn_math_fallback()
      self.doc.text(latex, code=True)
      return
    self._pin_inline_picture_offsets(run)
    # Shift down by the descent so formula baseline aligns with text baseline.
    # `<w:position>` is in half-points; positive = raised.
    descent_pt = baseline_mm / 25.4 * 72.0
    if descent_pt > 0.1:
      self._set_run_vertical_position(run, -round(descent_pt * 2))

  def _render_math_block_image(self, latex:str):
    """Centered block image for a display formula outside the OMML subset."""
    from .math import render_math_png
    from ..constants import Defaults, Align
    body_pt = self.doc._style.font_size or Defaults.FONT_SIZE
    # Display math is a touch larger than inline, matching LaTeX conventions.
    result = render_math_png(latex, fontsize_pt=body_pt * 1.25,
      color=self.style.body_color, fontset=self.style.math_fontset)
    if result is None:
      # Render the LaTeX body as centered code - never echo the `$$` delimiters.
      self._warn_math_fallback()
      self.doc.para(align=Align.CENTER)
      self.doc.text(latex, code=True)
      return
    buf, w_mm, h_mm, _baseline = result
    content_w = self.doc._page.content_width
    if w_mm > content_w:  # cap to page width
      h_mm = h_mm * content_w / w_mm
      w_mm = content_w
    self._insert_picture(buf, w_mm, h_mm)

  def _set_run_vertical_position(self, run, half_points:int):
    """Set `<w:position>` on a run (raises/lowers it; half-point units)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rpr = run._element.get_or_add_rPr()
    pos = rpr.find(qn("w:position"))
    if pos is None:
      pos = OxmlElement("w:position")
      rpr.append(pos)
    pos.set(qn("w:val"), str(half_points))

  def _warn_math_fallback(self):
    """Warn on total render failure. Message varies: matplotlib missing vs.
    present but unable to parse this LaTeX."""
    import warnings
    try:
      import matplotlib  # noqa: F401
      msg = ("math formula is outside both the OMML subset and matplotlib "
        "mathtext; rendering LaTeX source as code text")
    except ImportError:
      msg = ("math formula could not be rendered (install matplotlib for the "
        "image fallback: pip install docmarq[math]); rendering LaTeX as text")
    warnings.warn(msg, RuntimeWarning, stacklevel=2)

  #------------------------------------------------------------------------------------ Images

  @staticmethod
  def _standalone_image_src(inline_token:Token) -> str|None:
    """Return image `src` when `inline_token` is exactly one image (ignoring
    whitespace-only text children), else `None`. Determines block vs. inline
    image treatment.
    """
    children = inline_token.children or []
    seen_image = None
    for c in children:
      if c.type == "image":
        if seen_image is not None:
          return None
        seen_image = c
      elif c.type == "text":
        if c.content.strip():
          return None
      else:
        return None
    if seen_image is None:
      return None
    return get_attr(seen_image, "src")

  def _render_block_image(self, src:str, inline_token:Token):
    """Embed an image as a centered block paragraph. Falls back to italic
    alt text when the file is missing. Title DSL `![alt](src "key=val")`
    controls sizing and alignment (see `image_utils.parse_image_dsl`).
    """
    alt = ""
    title = None
    children = inline_token.children or []
    if children:
      img = next((c for c in children if c.type == "image"), None)
      if img is not None:
        alt = get_attr(img, "alt") or ""
        if not alt and img.children:
          alt = "".join(c.content for c in img.children if c.type == "text")
        title = get_attr(img, "title")
    dsl = image_utils.parse_image_dsl(title)
    path = image_utils.resolve_path(src, self.base_dir)
    if not path or not os.path.isfile(path):
      self.doc.para()
      self.doc.text(alt or f"[image: {src}]", italic=True)
      return
    if not self._try_insert_image(path, dsl=dsl):
      self.doc.para()
      self.doc.text(alt or f"[image: {src}]", italic=True)

  def _try_insert_image(self, path:str, dsl:image_utils.ImageDSL|None=None) -> bool:
    """Insert image scaled per DSL overrides with `style.image_max_h` cap.

    Every image goes through Pillow to:
      1. Crop transparent padding - Word renders the full transparent canvas,
         causing visible right-shift / shrunk content.
      2. Normalize format - JPEGs with APP2 ICC profile are rejected by
         python-docx; re-saving as PNG fixes it.
      3. Rasterize SVG (python-docx has no SVG support).

    `dsl.align` is applied to the paragraph after the insert.
    """
    from docx.image.exceptions import UnrecognizedImageError
    content_w = self.doc._page.content_width
    max_h = self.style.image_max_h
    dsl = dsl or image_utils.ImageDSL()
    buf = image_utils.preprocess_to_buffer(path)
    try:
      if buf is None:
        target_w, target_h = self._compute_dims(path, content_w, max_h, dsl)
        self._insert_picture(path, target_w, target_h)
      else:
        target_w, target_h = self._compute_dims_buffer(buf, content_w, max_h, dsl)
        buf.seek(0)
        self._insert_picture(buf, target_w, target_h)
    except (OSError, ValueError, UnrecognizedImageError):
      return False
    self._apply_image_align(dsl.align)
    return True

  @staticmethod
  def _compute_dims(path:str, content_w:float, max_h:float,
      dsl:image_utils.ImageDSL) -> tuple[float|None, float]:
    """Read natural dims from disk and apply DSL overrides."""
    try:
      from PIL import Image
      with Image.open(path) as im:
        nat_w, nat_h = im.size
    except (ImportError, OSError, ValueError):
      return (None, max_h)
    return image_utils.apply_dsl_dims(nat_w, nat_h, content_w, max_h, dsl)

  @staticmethod
  def _compute_dims_buffer(buf, content_w:float, max_h:float,
      dsl:image_utils.ImageDSL) -> tuple[float|None, float]:
    """Read natural dims from a buffer and apply DSL overrides."""
    try:
      from PIL import Image
      buf.seek(0)
      with Image.open(buf) as im:
        nat_w, nat_h = im.size
      buf.seek(0)
    except (ImportError, OSError, ValueError):
      return (None, max_h)
    return image_utils.apply_dsl_dims(nat_w, nat_h, content_w, max_h, dsl)

  def _apply_image_align(self, align:str|None):
    """Set alignment on the most recently added paragraph (the one that
    holds the just-inserted picture). No-op when `align` is unset."""
    if not align: return
    paragraphs = self.doc._doc.paragraphs
    if not paragraphs: return
    from ..utils import align_to_docx
    a = align_to_docx(align)
    if a is not None:
      paragraphs[-1].alignment = a

  #-------------------------------------------------------------------------------------- Mermaid

  def _render_mermaid(self, source:str, info_rest:str=""):
    """Render mermaid via `mmdc` (or `mermaid.ink`) to PNG and embed.
    `info_rest` is parsed as image DSL. Falls back to a code block on failure."""
    s = self.style
    if not s.mermaid_enable:
      self._fallback_mermaid_code(source)
      return
    png_path = mermaid.compile_to_png(source,
      cli=s.mermaid_cli, theme=s.mermaid_theme,
      background=s.mermaid_background, scale=s.mermaid_scale,
      font_family=s.body_family, font_dir=self.font_dir,
    )
    dsl = image_utils.parse_image_dsl(info_rest) if info_rest else None
    if png_path is None or not self._try_insert_image(png_path, dsl=dsl):
      self._fallback_mermaid_code(source)

  def _fallback_mermaid_code(self, source:str):
    """Render mermaid source as a regular fenced code block - used when
    `mmdc` is unavailable or the diagram fails to compile."""
    self.doc.code_block(
      source,
      language="mermaid",
      bg_color=self.style.code_block_bg,
      border_color=self.style.code_block_border,
      font_family=self.style.mono_family,
    )

  def _insert_picture(self, src, width_mm:float|None, height_mm:float|None):
    """Centered paragraph + picture run with explicit dimensions. `src` may
    be a path or `BytesIO`. Bypasses `doc.image()` to pass both width and
    height simultaneously after scaled dims are computed.
    """
    from docx.shared import Mm, Pt
    from ..utils import align_to_docx, to_mm
    self.doc._flush_para()
    p = self.doc._doc.add_paragraph()
    p.alignment = align_to_docx(Align.CENTER)
    pf = p.paragraph_format
    pf.left_indent = Mm(0)
    pf.right_indent = Mm(0)
    pf.first_line_indent = Mm(0)
    pf.space_before = Pt(self.doc._collapsed_before(2))
    pf.space_after = Pt(2)
    run = p.add_run()
    kwargs = {}
    if width_mm is not None:
      kwargs["width"] = Mm(to_mm(width_mm, self.doc.unit))
    if height_mm is not None:
      kwargs["height"] = Mm(to_mm(height_mm, self.doc.unit))
    run.add_picture(src, **kwargs)
    self._pin_inline_picture_offsets(run)
    self.doc._track_block_spacing(2)

  @staticmethod
  def _pin_inline_picture_offsets(run):
    """Set `distT/B/L/R="0"` and inject a zero `<wp:effectExtent>` on every
    `<wp:inline>` in this run. python-docx omits these; LibreOffice and some
    Word versions render a small horizontal offset when they are absent.
    """
    from docx.oxml import OxmlElement
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    def wp(tag):
      return f"{{{WP_NS}}}{tag}"
    for inline in run._element.findall(f".//{wp('inline')}"):
      for attr in ("distT", "distB", "distL", "distR"):
        inline.set(attr, "0")
      eff = inline.find(wp("effectExtent"))
      if eff is None:
        eff = OxmlElement("wp:effectExtent")
        for attr in ("l", "t", "r", "b"):
          eff.set(attr, "0")
        # OOXML schema order: extent → effectExtent → docPr → ...
        extent = inline.find(wp("extent"))
        if extent is not None:
          idx = list(inline).index(extent) + 1
          inline.insert(idx, eff)
        else:
          inline.insert(0, eff)

  #------------------------------------------------------------------------------- Footnote block

  def _render_footnote_block(self, tokens:list[Token], start:int) -> int:
    """Render the footnote bibliography section appended by markdown-it.
    Emits an H2 heading (when `style.footnote_label` is set) or a thin HR,
    then one paragraph per footnote prefixed with `[N]`. Runs are one pt
    smaller than body (`smaller_size`) so the bibliography reads as compact
    reference material. Heading keeps its style-driven size.
    """
    from ..constants import Defaults
    from ..utils import smaller_size
    end = find_close(tokens, start, "footnote_block_open", "footnote_block_close")
    if self.style.footnote_label:
      self.doc.heading(self.style.footnote_label, level=2)
    else:
      self.doc.hr(color=self.style.hr_color)
    body_pt = self.doc._style.font_size or Defaults.FONT_SIZE
    biblio_pt = smaller_size(body_pt)
    self.doc.font(size=biblio_pt)
    try:
      j = start + 1
      while j < end:
        if tokens[j].type == "footnote_open":
          item_end = find_close(tokens, j, "footnote_open", "footnote_close")
          label = (tokens[j].meta or {}).get("label") or "?"
          self._render_footnote_item(tokens, j + 1, item_end, label)
          j = item_end + 1
        else:
          j += 1
    finally:
      self.doc.font(size=body_pt)
    return end + 1

  def _render_footnote_item(self, tokens:list[Token], start:int, end:int,
      label:str):
    """Render a single footnote item: a paragraph beginning with `[label] `
    followed by the footnote's inline content. Multi-paragraph footnotes
    emit additional plain paragraphs (no further prefix).
    """
    first = True
    k = start
    while k < end:
      t = tokens[k]
      if t.type == "paragraph_open":
        inline = tokens[k+1]
        if first:
          self.doc.para()
          self.doc.text(f"[{label}] ", bold=True)
          self._add_runs_to_current_para(inline)
          first = False
        else:
          self._render_paragraph(inline)
        k += 3
      elif t.type == "footnote_anchor":
        k += 1  # back-reference decoration; no DOCX equivalent
      else:
        k += 1

  #------------------------------------------------------------------------------------ Table

  def _render_table(self, tokens:list[Token], start:int) -> int:
    end = find_close(tokens, start, "table_open", "table_close")
    header_cells: list[str] = []
    body_rows: list[list[str]] = []
    aligns: list[str] = []
    in_header = False
    current_row: list[str] = []
    j = start + 1
    while j < end:
      t = tokens[j]
      tt = t.type
      if tt == "thead_open": in_header = True
      elif tt == "thead_close": in_header = False
      elif tt == "tr_open": current_row = []
      elif tt == "tr_close":
        if in_header: header_cells = current_row
        else: body_rows.append(current_row)
      elif tt in ("th_open", "td_open"):
        if tt == "th_open":
          style_attr = get_attr(t, "style") or ""
          if "center" in style_attr: aligns.append(Align.CENTER)
          elif "right" in style_attr: aligns.append(Align.RIGHT)
          else: aligns.append(Align.LEFT)
        close_type = "th_close" if tt == "th_open" else "td_close"
        k = j + 1
        cell_text = ""
        while k < end and tokens[k].type != close_type:
          if tokens[k].type == "inline": cell_text = self._inline_to_plain(tokens[k])
          k += 1
        current_row.append(cell_text)
        j = k
      j += 1
    self.doc.table(
      header=header_cells if header_cells else None,
      body=body_rows,
      aligns=aligns if aligns else None,
    )
    return end + 1

  #------------------------------------------------------------------------------ Helpers

  @staticmethod
  def _inline_to_plain(inline_token:Token) -> str:
    """Flatten an inline token into a plain string for cells/short snippets."""
    parts = []
    for c in (inline_token.children or []):
      if c.type in ("text", "code_inline"):
        parts.append(c.content)
      elif c.type == "math_inline":
        # Tables take plain strings; preserve LaTeX in `$...$` notation.
        parts.append(f"${c.content}$")
      elif c.type == "softbreak": parts.append(" ")
      elif c.type == "hardbreak": parts.append("\n")
    return "".join(parts)

#---------------------------------------------------------------------------------- md_to_docx

def md_to_docx(
  md_text: str,
  output_path: str,
  style: MarkdownStyle|None = None,
  width: float|None = None,
  height: float|None = None,
  margin: float|tuple|None = None,
  metadata: dict|None = None,
  landscape: bool|None = None,
  base_dir: str|None = None,
  font_dir: str|None = None,
) -> DOCX:
  """Convert markdown text to a `.docx` file.

  YAML frontmatter `render:` sub-block controls page geometry, fonts,
  chrome, and locale. See `docmarq.md.render.RenderConfig`.

  Precedence: `MarkdownStyle()` defaults < `render.lang` preset < other
  `render:` keys < caller's `style=` non-default fields.

  Args:
    md_text: Markdown source.
    output_path: Destination `.docx` path.
    style: Optional `MarkdownStyle`. `None` uses GitHub-light defaults.
    width / height: Page dimensions in mm. `None` (default) reads
      `render.page` from frontmatter, falling back to A4.
    margin: Page margins in mm. `None` (default) reads `render.margin`,
      falling back to 20.
    metadata: Optional `dict` passed to `DOCX.metadata()`.
    landscape: Flip page. `None` (default) reads `render.landscape`.
      Top-level `landscape:` is no longer honored (warns).
    base_dir: Root for resolving relative image paths. `None` uses cwd.
    font_dir: Optional TTF root for mermaid diagram font sync. When set,
      diagrams render in `style.body_family` instead of system default.
  """
  from .render import (
    parse_render_block, build_style, warn_top_level_landscape,
  )
  fm = peek_frontmatter(md_text)
  warn_top_level_landscape(fm)
  render = parse_render_block(fm)
  if width is None:
    width = render.page.width if render.page else 210
  if height is None:
    height = render.page.height if render.page else 297
  if margin is None:
    margin = render.margin if render.margin is not None else 20
  if landscape is None:
    landscape = bool(render.landscape)
  if landscape:
    width, height = height, width
  eff_style = build_style(fm, style, render)
  doc = DOCX(output_path, width=width, height=height, margin=margin,
    gutter=render.gutter or 0)
  # Render-block typography that doesn't live on `MarkdownStyle` is applied
  # via the fluent doc API instead. Word picks these up for subsequent runs.
  # Typography from render block that isn't on MarkdownStyle goes via doc API.
  if render.font_size is not None:
    doc.font(size=render.font_size)
  if render.line_height is not None:
    doc.style(line_height=render.line_height)
  if metadata:
    doc.metadata(**metadata)
  renderer = MarkdownRenderer(doc, eff_style, base_dir=base_dir, font_dir=font_dir)
  renderer.render(md_text)
  doc.save()
  return doc
